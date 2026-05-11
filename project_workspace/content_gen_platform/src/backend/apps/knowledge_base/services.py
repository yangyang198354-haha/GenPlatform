"""Knowledge base processing: chunking, embedding, and RAG retrieval."""
import logging
import re
from datetime import timedelta
from pathlib import Path
from typing import TYPE_CHECKING, List, Optional

from django.conf import settings
from django.utils import timezone
from pgvector.django import L2Distance

from .models import Document, DocumentChunk

if TYPE_CHECKING:
    from sentence_transformers import SentenceTransformer

logger = logging.getLogger(__name__)

# A document is considered "orphaned" if it has been in status='processing'
# for longer than this threshold. Must be larger than the Celery hard time
# limit (_TIME_LIMIT = 660 s in tasks.py) so we never reap a still-running
# task — only ones whose worker was SIGKILL'd (typically by the OS OOM killer)
# before the SoftTimeLimitExceeded handler could write status='error'.
STALE_PROCESSING_THRESHOLD_SECONDS = 15 * 60  # 15 min (= 660 s hard kill + 240 s buffer)

# User-facing message written to error_message when a doc is reaped. Includes
# remediation guidance because the most common cause is server memory pressure
# and the user can retry once the system has more headroom.
STALE_PROCESSING_REAP_MESSAGE = (
    "处理被中断（worker 进程在加载向量模型时异常退出，通常是服务器内存不足导致）。"
    "请点击重试，或删除此文档后重新上传。"
)

# Lazy-loaded singleton embedding model
_embedding_model: "SentenceTransformer | None" = None


def _get_embedding_model() -> "SentenceTransformer":
    global _embedding_model
    if _embedding_model is None:
        from sentence_transformers import SentenceTransformer  # noqa: PLC0415
        logger.info("Loading embedding model: %s", settings.EMBEDDING_MODEL)
        _embedding_model = SentenceTransformer(
            settings.EMBEDDING_MODEL, device=settings.EMBEDDING_DEVICE
        )
    return _embedding_model


def _ocr_pdf(file_path: str) -> str:
    """
    OCR fallback for scanned PDFs (no embedded text layer).

    Converts each PDF page to a raster image via pdf2image (wraps poppler),
    then runs pytesseract to extract text.  Both libraries must be installed
    AND the corresponding system binaries must be present:
      - poppler-utils  (for pdf2image / pdftoppm)
      - tesseract-ocr + tesseract-ocr-chi-sim + tesseract-ocr-eng
        (for pytesseract)

    If either library is missing the function logs a warning and returns ""
    so the caller can surface a meaningful empty result rather than crashing.
    """
    try:
        from pdf2image import convert_from_path   # noqa: PLC0415
        import pytesseract                        # noqa: PLC0415
    except ImportError as exc:
        logger.warning(
            "OCR unavailable for scanned PDF '%s': %s — "
            "install pdf2image, pytesseract, and system packages "
            "poppler-utils + tesseract-ocr + tesseract-ocr-chi-sim",
            file_path, exc,
        )
        return ""

    try:
        # DPI=200: good balance between accuracy and memory for A4 pages.
        # lang="chi_sim+eng": mixed Chinese/English documents are common.
        images = convert_from_path(file_path, dpi=200)
        texts = [
            pytesseract.image_to_string(img, lang="chi_sim+eng")
            for img in images
        ]
        return "\n".join(t for t in texts if t.strip())
    except Exception as exc:
        raise RuntimeError(f"PDF OCR failed: {exc}") from exc


def _extract_text(file_path: str, file_type: str) -> str:
    """Extract plain text from a document file.

    For PDF files the extraction is two-stage:
      1. Try pypdf's text-layer extraction (fast, zero extra dependencies).
      2. If every page returns empty text (i.e. the file is a scanned /
         image-only PDF), fall back to OCR via pdf2image + pytesseract.
    """
    path = Path(file_path)
    if file_type == "txt" or file_type == "md":
        return path.read_text(encoding="utf-8", errors="replace")

    if file_type == "pdf":
        try:
            import pypdf  # noqa: PLC0415
            reader = pypdf.PdfReader(str(path))
            pages_text = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(pages_text)

            # If all pages have no text it is almost certainly a scanned PDF.
            # Fall back to OCR so users get content instead of an empty chunk.
            if not text.strip():
                logger.info(
                    "PDF '%s' has no embedded text layer — falling back to OCR",
                    file_path,
                )
                text = _ocr_pdf(file_path)

            return text
        except Exception as e:
            raise RuntimeError(f"PDF extraction failed: {e}") from e

    if file_type == "docx":
        try:
            import docx  # noqa: PLC0415
            doc = docx.Document(str(path))
            texts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            texts.append(cell.text)
            return "\n".join(texts)
        except Exception as e:
            raise RuntimeError(f"DOCX extraction failed: {e}") from e

    raise ValueError(f"Unsupported file type: {file_type}")


def _chunk_text(text: str, chunk_size: int, overlap: int) -> List[str]:
    """Split text into overlapping chunks by character count."""
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return [c for c in chunks if c.strip()]


def reap_stale_processing_documents(user_id: Optional[int] = None) -> int:
    """Mark orphaned 'processing' documents as 'error' so they stop displaying
    a phantom progress bar forever.

    A document is "orphaned" when its Celery worker was SIGKILL'd (most
    commonly by the OS OOM killer during embedding-model load) before any
    Python exception handler could write status='error'. The soft/hard
    time-limit safety net only fires when the worker process is still alive;
    a hard SIGKILL silently leaves the row at status='processing' forever.

    This function is the only way orphaned rows recover. It is called from:
      - the Celery worker_ready signal (on every worker startup → cleans up
        whatever was running when the previous worker died)
      - the document list view (so users see recovery the moment they refresh
        the KB page — no need to wait for a restart)

    The threshold (STALE_PROCESSING_THRESHOLD_SECONDS) is deliberately larger
    than the task's hard time_limit so a legitimately long-running task is
    never misclassified as orphaned.

    Args:
        user_id: when given, only reap that user's documents (used from the
            list view to avoid scanning the global table on every request).
            When None, sweep across all users (used at worker startup).

    Returns:
        The number of documents transitioned from 'processing' to 'error'.
    """
    cutoff = timezone.now() - timedelta(seconds=STALE_PROCESSING_THRESHOLD_SECONDS)
    qs = Document.objects.filter(status="processing", created_at__lt=cutoff)
    if user_id is not None:
        qs = qs.filter(user_id=user_id)

    reaped_ids = list(qs.values_list("pk", flat=True))
    if not reaped_ids:
        return 0

    n = Document.objects.filter(pk__in=reaped_ids).update(
        status="error",
        error_message=STALE_PROCESSING_REAP_MESSAGE,
        progress_message="处理被中断",
    )
    logger.warning(
        "Reaped %d orphaned 'processing' documents (ids=%s, user_id=%s)",
        n,
        reaped_ids,
        user_id,
    )
    return n


def _update_progress(document_id: int, progress: int, message: str) -> None:
    """Update document progress fields in-place.

    Called at each major processing stage so the frontend can show a
    progress bar instead of a static "处理中" tag.
    Only updates the two progress fields — does NOT touch status.
    """
    Document.objects.filter(pk=document_id).update(
        progress=progress,
        progress_message=message,
    )


def process_document(document_id: int) -> None:
    """
    Celery task body: extract text → chunk → embed → store in pgvector.
    Marks document status as 'available' on success, 'error' on failure.

    Progress stages:
      5  — task started, resetting state
      20 — text extraction done
      40 — text chunking done
      60 — embedding model loaded
      90 — embeddings computed, writing to DB
     100 — finished (status → available)
    """
    try:
        doc = Document.objects.get(pk=document_id)
        doc.status = "processing"
        doc.progress = 5
        doc.progress_message = "任务已开始"
        doc.save(update_fields=["status", "progress", "progress_message"])

        _update_progress(document_id, 10, "正在提取文本")
        text = _extract_text(doc.file_path, doc.file_type)
        _update_progress(document_id, 20, "文本提取完成，正在切片")

        chunks = _chunk_text(text, settings.CHUNK_SIZE, settings.CHUNK_OVERLAP)
        _update_progress(document_id, 40, f"切片完成（共 {len(chunks)} 个），正在加载向量模型")

        model = _get_embedding_model()
        _update_progress(document_id, 60, f"正在计算向量嵌入（共 {len(chunks)} 个切片）")

        # batch_size=32 caps peak RAM: encoding all chunks at once can OOM
        # the worker when processing large documents (especially DOCX/PDF).
        embeddings = model.encode(
            chunks, batch_size=32, show_progress_bar=False, normalize_embeddings=True
        )
        _update_progress(document_id, 90, "向量计算完成，正在写入数据库")

        # Delete existing chunks (re-processing case)
        DocumentChunk.objects.filter(document=doc).delete()

        chunk_objects = [
            DocumentChunk(
                document=doc,
                chunk_index=i,
                content=chunk,
                embedding=embedding.tolist(),
            )
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings))
        ]
        DocumentChunk.objects.bulk_create(chunk_objects)

        doc.refresh_from_db()
        doc.status = "available"
        doc.chunk_count = len(chunk_objects)
        doc.error_message = ""
        doc.progress = 100
        doc.progress_message = "处理完成"
        doc.save(update_fields=["status", "chunk_count", "error_message", "progress", "progress_message"])
        logger.info("Document %d processed: %d chunks", document_id, len(chunk_objects))

    except Document.DoesNotExist:
        logger.error("Document %d not found", document_id)
    except Exception as e:
        logger.exception("Document %d processing failed", document_id)
        Document.objects.filter(pk=document_id).update(
            status="error",
            error_message=str(e),
            progress_message="处理失败",
        )


def _extract_cn_keywords(text: str) -> List[str]:
    """
    Extract candidate Chinese keyword phrases from text without external deps.

    Scans all runs of 2+ consecutive Chinese characters and generates sliding
    windows of length 2-4.  Longer windows are returned first so callers can
    prefer more-specific terms when building OR-filter chains.

    Example:
        "业主大会中车位是否计入投票权数"
        → ["业主大会", "主大会中", ..., "车位", "投票权数", ...]
    """
    keywords: list[str] = []
    seen: set[str] = set()
    # Find all runs of 2+ consecutive CJK unified ideographs
    for run in re.findall(r'[\u4e00-\u9fff]{2,}', text):
        for n in (4, 3, 2):           # longer first → more specific
            for i in range(len(run) - n + 1):
                kw = run[i : i + n]
                if kw not in seen:
                    seen.add(kw)
                    keywords.append(kw)
    return keywords


def search(user_id: int, query: str, top_k: int = 5) -> List[DocumentChunk]:
    """
    Hybrid RAG retrieval: semantic + anchor-chunk injection + keyword fallback.

    Pure semantic (vector) search fails for queries like "tell me the community
    name" when the name only appears in the document header (chunk_index=0) and
    the query is semantically about a different topic (e.g. voting rules).

    Three-layer strategy
    --------------------
    1. **Semantic** — L2Distance on embeddings finds the thematically closest
       chunks (e.g. voting-rule paragraphs).  Fetches top_k * 2 candidates to
       give the later merge stages more material to work with.

    2. **Anchor injection** — chunk_index=0 of every document that produced a
       semantic hit is always included.  The first chunk of a document typically
       contains its title and key named entities (e.g. the community name
       "红枫岭枫和苑").  This ensures document-level metadata is never silently
       dropped regardless of query topic.

    3. **Keyword fallback** — runs a simple icontains OR-filter for 2-4 char
       Chinese ngrams extracted from the query.  Catches chunks that mention
       query terms verbatim but score poorly on cosine/L2 distance (e.g. a
       definitions section that lists key terms).

    Returns at most top_k + (number of injected anchor chunks) results so that
    anchors never displace genuinely relevant semantic hits.
    """
    from django.db.models import Q  # noqa: PLC0415

    model = _get_embedding_model()
    query_embedding = model.encode([query], normalize_embeddings=True)[0].tolist()

    base_qs = DocumentChunk.objects.filter(
        document__user_id=user_id,
        document__status="available",
    )

    # ── Layer 1: Semantic search ─────────────────────────────────────────────
    semantic_results: List[DocumentChunk] = list(
        base_qs
        .annotate(distance=L2Distance("embedding", query_embedding))
        .order_by("distance")[: top_k * 2]
    )
    seen_pks: set[int] = {c.pk for c in semantic_results}
    hit_doc_ids: set[int] = {c.document_id for c in semantic_results}

    # ── Layer 2: Anchor chunk injection ──────────────────────────────────────
    # Always include chunk_index=0 of every document that had a semantic hit.
    # "Natural top_k" = the first top_k semantic candidates (what pure semantic
    # would return).  Anchors already inside the natural top_k are naturally
    # included and do NOT consume an extra slot.  Only anchors that fall *outside*
    # the natural top_k are force-injected as bonus entries.
    natural_top_k_pks: set[int] = {c.pk for c in semantic_results[:top_k]}
    all_anchor_chunks: List[DocumentChunk] = list(
        base_qs.filter(document_id__in=hit_doc_ids, chunk_index=0)
    )
    # Extra anchors: not already in the natural top_k (need a free slot)
    extra_anchor_chunks: List[DocumentChunk] = [
        c for c in all_anchor_chunks if c.pk not in natural_top_k_pks
    ]
    seen_pks.update(c.pk for c in extra_anchor_chunks)

    # ── Layer 3: Keyword fallback ─────────────────────────────────────────────
    keywords = _extract_cn_keywords(query)
    keyword_chunks: List[DocumentChunk] = []
    if keywords:
        kw_q = Q()
        for kw in keywords[:15]:     # cap OR-chain length to avoid slow plans
            kw_q |= Q(content__icontains=kw)
        keyword_chunks = list(
            base_qs
            .filter(kw_q)
            .exclude(pk__in=seen_pks)
            .order_by("chunk_index")[: top_k]
        )

    # ── Merge ────────────────────────────────────────────────────────────────
    # natural top_k semantic results (hard cap = top_k)
    # + extra anchor chunks that fell outside natural top_k (force-included)
    # + keyword extras (deduped)
    natural_results = semantic_results[:top_k]
    return natural_results + extra_anchor_chunks + keyword_chunks
